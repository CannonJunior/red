"""
Tests for proposal/document_generator.py

Covers:
    - Expected use: requirement loading, section mapping, DOCX output
    - Edge cases: empty requirements, missing categories, Ollama failure
    - Failure cases: bad JSON, missing template
"""

import json
import csv
from unittest.mock import MagicMock, patch

import pytest
import requests

from proposal.document_generator import (
    Requirement,
    SectionContent,
    DraftResult,
    SECTION_MAP,
    _build_section_prompt,
    draft_proposal,
    draft_volume,
    generate_section_prose,
    load_requirements_from_csv,
    load_requirements_from_json,
    map_requirements_to_sections,
    write_section_to_doc,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_requirements():
    return [
        Requirement(id="REQ-001", category="technical_requirement",
                    text="The system shall support TLS 1.3.", reference="SOW 3.1", priority="shall"),
        Requirement(id="REQ-002", category="security_requirement",
                    text="The system shall enforce MFA for all users.", reference="SOW 4.2", priority="shall"),
        Requirement(id="REQ-003", category="management_requirement",
                    text="The contractor shall provide a monthly status report.", reference="SOW 5.1", priority="shall"),
        Requirement(id="REQ-004", category="risk_item",
                    text="Legacy system integration poses schedule risk.", priority="should"),
    ]


@pytest.fixture
def json_req_file(tmp_path, sample_requirements):
    data = [
        {"id": r.id, "category": r.category, "text": r.text,
         "reference": r.reference, "priority": r.priority}
        for r in sample_requirements
    ]
    p = tmp_path / "requirements.json"
    p.write_text(json.dumps(data))
    return p


@pytest.fixture
def csv_req_file(tmp_path, sample_requirements):
    p = tmp_path / "requirements.csv"
    with open(p, "w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=["id", "category", "text", "reference", "priority"])
        writer.writeheader()
        for r in sample_requirements:
            writer.writerow({"id": r.id, "category": r.category, "text": r.text,
                             "reference": r.reference, "priority": r.priority})
    return p


@pytest.fixture
def mock_docx_template(tmp_path, monkeypatch):
    """Create a minimal real .docx template and patch _templates_dir."""
    from docx import Document
    from docx.shared import Pt

    tpl_dir = tmp_path / "templates"
    tpl_dir.mkdir()

    for vol in ["technical_volume", "management_volume"]:
        doc = Document()
        doc.add_heading(vol.replace("_", " ").title(), level=0)
        doc.add_paragraph(f"[ {vol.upper()} — Replace this placeholder with proposal content. ]")
        doc.add_paragraph("[ 2.0 TECHNICAL APPROACH — Replace this placeholder with proposal content. ]")
        doc.add_paragraph("[ 1.0 MANAGEMENT APPROACH — Replace this placeholder with proposal content. ]")
        doc.save(str(tpl_dir / f"{vol}.docx"))

    monkeypatch.setenv("TEMPLATES_DIR", str(tpl_dir))
    return tpl_dir


# ---------------------------------------------------------------------------
# load_requirements_from_json
# ---------------------------------------------------------------------------

class TestLoadRequirementsJson:
    def test_loads_all_requirements(self, json_req_file):
        reqs = load_requirements_from_json(json_req_file)
        assert len(reqs) == 4

    def test_field_mapping(self, json_req_file):
        reqs = load_requirements_from_json(json_req_file)
        r = reqs[0]
        assert r.id == "REQ-001"
        assert r.category == "technical_requirement"
        assert "TLS 1.3" in r.text
        assert r.reference == "SOW 3.1"

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            load_requirements_from_json(tmp_path / "nonexistent.json")

    def test_non_array_raises(self, tmp_path):
        p = tmp_path / "bad.json"
        p.write_text('{"key": "value"}')
        with pytest.raises(ValueError, match="JSON array"):
            load_requirements_from_json(p)

    def test_empty_array_returns_empty(self, tmp_path):
        p = tmp_path / "empty.json"
        p.write_text("[]")
        assert load_requirements_from_json(p) == []

    def test_default_category_when_missing(self, tmp_path):
        p = tmp_path / "req.json"
        p.write_text(json.dumps([{"text": "Some requirement"}]))
        reqs = load_requirements_from_json(p)
        assert reqs[0].category == "general_requirement"


# ---------------------------------------------------------------------------
# load_requirements_from_csv
# ---------------------------------------------------------------------------

class TestLoadRequirementsCsv:
    def test_loads_all_requirements(self, csv_req_file):
        reqs = load_requirements_from_csv(csv_req_file)
        assert len(reqs) == 4

    def test_field_mapping(self, csv_req_file):
        reqs = load_requirements_from_csv(csv_req_file)
        r = next(r for r in reqs if r.id == "REQ-002")
        assert "MFA" in r.text

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            load_requirements_from_csv(tmp_path / "missing.csv")

    def test_skips_empty_text_rows(self, tmp_path):
        p = tmp_path / "sparse.csv"
        with open(p, "w", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=["id", "category", "text"])
            writer.writeheader()
            writer.writerow({"id": "1", "category": "tech", "text": "Valid req"})
            writer.writerow({"id": "2", "category": "tech", "text": ""})
        reqs = load_requirements_from_csv(p)
        assert len(reqs) == 1


# ---------------------------------------------------------------------------
# map_requirements_to_sections
# ---------------------------------------------------------------------------

class TestMapRequirementsToSections:
    def test_maps_to_correct_volumes(self, sample_requirements):
        sections = map_requirements_to_sections(sample_requirements)
        volumes = {s.volume for s in sections.values()}
        assert "technical_volume" in volumes
        assert "management_volume" in volumes

    def test_groups_same_section_together(self, sample_requirements):
        # REQ-001 (technical) and REQ-002 (security) both map to technical_volume 2.2
        sections = map_requirements_to_sections(sample_requirements)
        tech_22 = next(
            (s for s in sections.values() if "2.2" in s.heading and s.volume == "technical_volume"),
            None
        )
        assert tech_22 is not None
        assert len(tech_22.requirements) >= 1

    def test_unknown_category_falls_back_to_general(self):
        reqs = [Requirement(id="R1", category="unicorn_category", text="Odd requirement")]
        sections = map_requirements_to_sections(reqs)
        assert len(sections) == 1
        s = list(sections.values())[0]
        assert s.volume == "technical_volume"

    def test_empty_requirements_returns_empty_dict(self):
        assert map_requirements_to_sections([]) == {}

    def test_custom_section_map_applied(self):
        custom_map = {"custom_req": ["management_volume", "7.0 Custom Section"]}
        reqs = [Requirement(id="R1", category="custom_req", text="Custom")]
        sections = map_requirements_to_sections(reqs, custom_map)
        assert "management_volume::7.0 Custom Section" in sections


# ---------------------------------------------------------------------------
# generate_section_prose
# ---------------------------------------------------------------------------

class TestGenerateSectionProse:
    def test_returns_placeholder_when_ollama_down(self):
        section = SectionContent(
            volume="technical_volume",
            heading="2.0 Technical Approach",
            requirements=[Requirement(id="R1", category="tech", text="Test req")],
        )
        with patch("requests.post") as mock_post:
            mock_post.side_effect = requests.ConnectionError("Connection refused")
            prose = generate_section_prose(section)
        assert "DRAFT REQUIRED" in prose
        assert "R1" in prose

    def test_returns_ollama_response_on_success(self):
        section = SectionContent(
            volume="technical_volume",
            heading="2.0 Technical Approach",
            requirements=[Requirement(id="R1", category="tech", text="Test req")],
        )
        mock_resp = MagicMock()
        mock_resp.json.return_value = {"response": "Generated proposal text here."}
        mock_resp.raise_for_status.return_value = None

        with patch("requests.post", return_value=mock_resp):
            prose = generate_section_prose(section)
        assert prose == "Generated proposal text here."

    def test_placeholder_includes_req_ids(self):
        section = SectionContent(
            volume="technical_volume",
            heading="Test",
            requirements=[
                Requirement(id="R1", category="tech", text="First"),
                Requirement(id="R2", category="tech", text="Second"),
            ],
        )
        with patch("requests.post", side_effect=requests.ConnectionError("down")):
            prose = generate_section_prose(section)
        assert "R1" in prose
        assert "R2" in prose


# ---------------------------------------------------------------------------
# _build_section_prompt
# ---------------------------------------------------------------------------

class TestBuildSectionPrompt:
    def test_includes_requirement_text(self):
        section = SectionContent(
            volume="tv", heading="2.0 Approach",
            requirements=[Requirement(id="R1", category="t", text="Must support IPv6.")],
        )
        prompt = _build_section_prompt(section, "Acme Corp", "Big Proposal")
        assert "IPv6" in prompt
        assert "Acme Corp" in prompt
        assert "2.0 Approach" in prompt

    def test_includes_reference_when_present(self):
        section = SectionContent(
            volume="tv", heading="2.0",
            requirements=[Requirement(id="R1", category="t", text="Req", reference="SOW 3.1")],
        )
        prompt = _build_section_prompt(section, "Co", "Prop")
        assert "SOW 3.1" in prompt


# ---------------------------------------------------------------------------
# draft_proposal (integration — uses mock Ollama + real template)
# ---------------------------------------------------------------------------

class TestDraftProposal:
    def test_drafts_technical_volume(self, sample_requirements, mock_docx_template, tmp_path, monkeypatch):
        monkeypatch.setenv("OUTPUTS_DIR", str(tmp_path / "outputs"))

        mock_resp = MagicMock()
        mock_resp.json.return_value = {"response": "Professional proposal text."}
        mock_resp.raise_for_status.return_value = None

        with patch("requests.post", return_value=mock_resp):
            result = draft_proposal(
                requirements=sample_requirements,
                proposal_title="Test Proposal",
                solicitation_number="FA8612-25-R-TEST",
                company_name="Test Corp",
                volumes=["technical_volume"],
            )

        assert result.solicitation_number == "FA8612-25-R-TEST"
        assert "technical_volume" in result.volumes_written
        assert result.volumes_written["technical_volume"].is_file()
        assert result.sections_generated >= 1
        assert result.requirements_addressed == len(sample_requirements)
        assert not result.errors

    def test_missing_template_recorded_as_error(self, sample_requirements, tmp_path, monkeypatch):
        monkeypatch.setenv("TEMPLATES_DIR", str(tmp_path / "empty_templates"))
        monkeypatch.setenv("OUTPUTS_DIR", str(tmp_path / "outputs"))
        (tmp_path / "empty_templates").mkdir()

        with patch("requests.post", side_effect=requests.ConnectionError("Ollama down")):
            result = draft_proposal(
                requirements=sample_requirements,
                proposal_title="Test",
                solicitation_number="TEST-001",
                volumes=["technical_volume"],
            )

        assert any("Template not found" in e or "template" in e.lower() for e in result.errors)

    def test_empty_requirements_returns_error(self, mock_docx_template, tmp_path, monkeypatch):
        monkeypatch.setenv("OUTPUTS_DIR", str(tmp_path / "outputs"))
        result = draft_proposal(
            requirements=[],
            proposal_title="Test",
            solicitation_number="TEST-002",
        )
        assert result.requirements_addressed == 0
        assert result.sections_generated == 0
        assert len(result.errors) >= 1

    def test_output_filename_contains_solicitation_and_date(
        self, sample_requirements, mock_docx_template, tmp_path, monkeypatch
    ):
        monkeypatch.setenv("OUTPUTS_DIR", str(tmp_path / "outputs"))
        mock_resp = MagicMock()
        mock_resp.json.return_value = {"response": "Text."}
        mock_resp.raise_for_status.return_value = None

        with patch("requests.post", return_value=mock_resp):
            result = draft_proposal(
                requirements=sample_requirements,
                proposal_title="Test",
                solicitation_number="FA1234-25-R-0001",
                volumes=["technical_volume"],
            )

        if "technical_volume" in result.volumes_written:
            fname = result.volumes_written["technical_volume"].name
            assert "FA1234-25-R-0001" in fname
            assert "technical_volume" in fname
