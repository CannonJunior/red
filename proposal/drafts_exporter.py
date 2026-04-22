"""
proposal/drafts_exporter.py — Export a proposal document to .docx.

Walks the section tree stored in proposal_documents / proposal_sections
and renders it as a Word document using python-docx.

Section nesting depth drives the heading level:
  depth 0 → Heading 1
  depth 1 → Heading 2
  depth 2 → Heading 3
  depth 3+ → Heading 4

Usage:
    from proposal.drafts_exporter import export_document_docx
    docx_bytes = export_document_docx(doc_id)
"""

from __future__ import annotations

import io
import logging
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)


def export_document_docx(doc_id: str) -> Optional[bytes]:
    """
    Export a proposal document and all its sections to a .docx file.

    Args:
        doc_id: proposal_documents.id to export.

    Returns:
        bytes: Raw .docx file content, or None if the document was not found.

    Raises:
        ImportError: If python-docx is not installed.
    """
    from docx import Document  # type: ignore[import]
    from docx.shared import Pt, RGBColor  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]

    from proposal.drafts_db import get_document

    doc_record = get_document(doc_id, include_sections=True)
    if doc_record is None:
        return None

    word_doc = Document()

    # Document title
    title_para = word_doc.add_heading(doc_record["title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle line: doc_type + status
    sub = word_doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"{doc_record.get('doc_type', '').replace('_', ' ').title()} — {doc_record.get('status', 'draft').title()}"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    word_doc.add_paragraph()  # spacer

    # Walk section tree
    _write_sections(word_doc, doc_record.get("sections") or [], depth=0)

    buf = io.BytesIO()
    word_doc.save(buf)
    return buf.getvalue()


def _write_sections(word_doc, sections: List[Dict], depth: int) -> None:
    """
    Recursively write sections into the Word document.

    Args:
        word_doc: python-docx Document object.
        sections: List of section dicts (with optional 'children' key).
        depth: Nesting depth (0 = top-level).
    """
    from docx.shared import RGBColor  # type: ignore[import]

    heading_level = min(depth + 1, 4)  # docx supports 1–9; cap at 4 for readability

    for sec in sections:
        # Section heading
        number = sec.get("section_number", "").strip()
        title = sec.get("title", "").strip()
        heading_text = f"{number}  {title}" if number else title
        word_doc.add_heading(heading_text, level=heading_level)

        # Guidance block (shown in amber in the UI; use italic grey in docx)
        guidance = sec.get("guidance", "").strip()
        if guidance:
            para = word_doc.add_paragraph()
            run = para.add_run(f"[Guidance] {guidance}")
            run.italic = True
            run.font.color.rgb = RGBColor(0xAA, 0x88, 0x44)

        # Section content
        content = (sec.get("content") or "").strip()
        if content:
            for line in content.splitlines():
                word_doc.add_paragraph(line) if line.strip() else word_doc.add_paragraph()
        else:
            para = word_doc.add_paragraph()
            run = para.add_run("[No content]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Owner annotation
        owner = (sec.get("owner") or "").strip()
        if owner:
            para = word_doc.add_paragraph()
            run = para.add_run(f"Owner: {owner}")
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x99, 0xCC)

        # Recurse into children
        children = sec.get("children") or []
        if children:
            _write_sections(word_doc, children, depth + 1)
