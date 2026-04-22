"""
proposal/color_team.py — Color team review ingestion for proposal documents.

Pipeline:
  1. Extract comments from a reviewer's marked-up .docx file
  2. Match each comment to the proposal section it annotates
  3. Detect which RFP compliance requirement the comment addresses
  4. Rate the section's compliance on the government five-point scale
  5. Persist results as proposal_comments with rating metadata

Rating scale (FAR 15.305 / DFARS 215.305 language):
  Outstanding   — Exceptional approach; significant strengths; no weaknesses
  Good          — Thorough approach; strengths outweigh weaknesses
  Acceptable    — Meets requirements; strengths and weaknesses offset
  Marginal      — Demonstrated deficiencies; correctable with revisions
  Unacceptable  — Fails to meet requirements; not awardable as written
"""

from __future__ import annotations

import io
import logging
import re
import zipfile
from typing import Dict, List, Optional
from xml.etree import ElementTree as ET

from proposal.database import get_conn
from proposal.drafts_db import (
    create_comment,
    _build_section_tree,
)

logger = logging.getLogger(__name__)

try:
    from config.model_config import get_model, ModelTier
    _DEFAULT_LOCAL_MODEL = get_model(ModelTier.LOCAL)
except Exception:
    _DEFAULT_LOCAL_MODEL = "qwen2.5:3b"

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

RATING_SCALE = ["Unacceptable", "Marginal", "Acceptable", "Good", "Outstanding"]

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


# ---------------------------------------------------------------------------
# Schema migration — adds rating/requirement columns to proposal_comments
# ---------------------------------------------------------------------------

_COLOR_TEAM_DDL = [
    "ALTER TABLE proposal_comments ADD COLUMN rating TEXT",
    "ALTER TABLE proposal_comments ADD COLUMN requirement_id TEXT",
    "ALTER TABLE proposal_comments ADD COLUMN requirement_text TEXT",
]


def run_color_team_migration() -> None:
    """
    Add color-team columns to proposal_comments if not present.

    Safe to call on every startup — each ALTER is wrapped in try/except.
    """
    with get_conn() as conn:
        for stmt in _COLOR_TEAM_DDL:
            try:
                conn.execute(stmt)
            except Exception:
                pass  # Column already exists


# ---------------------------------------------------------------------------
# Word document comment extraction
# ---------------------------------------------------------------------------

def extract_docx_comments(docx_bytes: bytes) -> List[Dict]:
    """
    Extract all reviewer comments from a .docx file.

    Reads the raw Open XML zip, parses word/comments.xml, and correlates
    each comment with the anchor text from word/document.xml.

    Args:
        docx_bytes: Raw bytes of the uploaded .docx file.

    Returns:
        list[dict]: Each dict has keys:
            id, author, date, content, anchor_text
        Empty list if the file contains no comments or is unreadable.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            try:
                comments_xml = z.read("word/comments.xml")
            except KeyError:
                logger.info("No word/comments.xml in docx — no comments.")
                return []
            doc_xml = z.read("word/document.xml")
    except Exception as exc:
        logger.error("Failed to open docx: %s", exc)
        return []

    # Unregister any pre-existing prefix to avoid ParseError in stdlib ET
    ET.register_namespace("w", W_NS)

    try:
        comments_root = ET.fromstring(comments_xml)
        doc_root = ET.fromstring(doc_xml)
    except ET.ParseError as exc:
        logger.error("XML parse error: %s", exc)
        return []

    anchor_map = _build_anchor_map(doc_root)

    results = []
    for comment in comments_root.iter(_w("comment")):
        cid = comment.get(_w("id"), "")
        author = comment.get(_w("author"), "")
        date = comment.get(_w("date"), "")

        # Collect all text runs inside the comment
        text_parts = [t.text for t in comment.iter(_w("t")) if t.text]
        content = "".join(text_parts).strip()
        if not content:
            continue

        results.append({
            "id": cid,
            "author": author,
            "date": date,
            "content": content,
            "anchor_text": anchor_map.get(cid, ""),
        })

    return results


def _build_anchor_map(doc_root: ET.Element) -> Dict[str, str]:
    """
    Walk document XML to collect the text ranges bracketed by each comment.

    Relies on w:commentRangeStart / w:commentRangeEnd pairs in the document.

    Args:
        doc_root: Parsed root element of word/document.xml.

    Returns:
        dict: Maps comment id string → anchor text string.
    """
    open_ranges: Dict[str, List[str]] = {}
    anchor_map: Dict[str, str] = {}

    for elem in doc_root.iter():
        tag = elem.tag
        if tag == _w("commentRangeStart"):
            cid = elem.get(_w("id"), "")
            if cid:
                open_ranges[cid] = []
        elif tag == _w("t") and elem.text:
            for parts in open_ranges.values():
                parts.append(elem.text)
        elif tag == _w("commentRangeEnd"):
            cid = elem.get(_w("id"), "")
            if cid in open_ranges:
                anchor_map[cid] = "".join(open_ranges.pop(cid)).strip()

    return anchor_map


# ---------------------------------------------------------------------------
# Section matching
# ---------------------------------------------------------------------------

_WORD_RE = re.compile(r"\w{4,}")


def build_section_index(sections: List[Dict]) -> Dict[str, frozenset]:
    """
    Pre-compute a keyword word-set for every section.

    Call this once per flat section list (not per comment) and pass the
    result to match_comment_to_section via the ``section_index`` parameter.

    Args:
        sections: Flat list of section dicts (each has 'id', 'content', 'title').

    Returns:
        dict: Maps section id → frozenset of lowercase words (length ≥ 4).
    """
    index: Dict[str, frozenset] = {}
    for sec in sections:
        haystack = f"{sec.get('title', '')} {sec.get('content', '')}".lower()
        index[sec["id"]] = frozenset(_WORD_RE.findall(haystack))
    return index


def match_comment_to_section(
    anchor_text: str,
    comment_text: str,
    sections: List[Dict],
    section_index: Optional[Dict[str, frozenset]] = None,
) -> Optional[str]:
    """
    Find the section ID whose content contains the comment's anchor text.

    Falls back to keyword overlap between comment text and section
    title/content when no anchor match is found.

    Args:
        anchor_text: The passage of text the comment is attached to.
        comment_text: The body of the reviewer's comment.
        sections: Flat list of section dicts (each has 'id', 'content', 'title').
        section_index: Optional pre-computed word-set map from build_section_index().
                       When provided, avoids re-running regex on every call.

    Returns:
        str or None: Section ID of the best match.
    """
    anchor_lower = anchor_text.lower() if anchor_text else ""

    # Pass 1: exact anchor substring match against lowercased content
    if anchor_lower:
        for sec in sections:
            if anchor_lower in (sec.get("content") or "").lower():
                return sec["id"]

    # Pass 2: keyword overlap — reuse pre-computed index when available
    comment_words = frozenset(_WORD_RE.findall(comment_text.lower()))
    best_id: Optional[str] = None
    best_score = 0
    for sec in sections:
        if section_index is not None:
            sec_words = section_index.get(sec["id"], frozenset())
        else:
            haystack = f"{sec.get('title', '')} {sec.get('content', '')}".lower()
            sec_words = frozenset(_WORD_RE.findall(haystack))
        overlap = len(comment_words & sec_words)
        if overlap > best_score:
            best_score = overlap
            best_id = sec["id"]

    return best_id if best_score >= 2 else None


# ---------------------------------------------------------------------------
# Compliance requirement loading
# ---------------------------------------------------------------------------

def fetch_shredding_requirements(opportunity_id: str) -> List[Dict]:
    """
    Load RFP requirements from the shredding output for an opportunity.

    Uses get_conn() so the call participates in the project's connection pool
    and WAL mode rather than opening a raw sqlite3 connection.

    Args:
        opportunity_id: Opportunity UUID.

    Returns:
        list[dict]: Each dict has 'id', 'section', 'source_text',
                    'compliance_type', 'requirement_category'.
    """
    try:
        with get_conn() as conn:
            rows = conn.execute(
                """SELECT id, section, source_text, compliance_type,
                          requirement_category, priority
                   FROM requirements
                   WHERE opportunity_id = ?
                   ORDER BY section, priority DESC""",
                (opportunity_id,),
            ).fetchall()
        return [dict(r) for r in rows]
    except Exception as exc:
        logger.warning("Could not load requirements for %s: %s", opportunity_id, exc)
        return []


# ---------------------------------------------------------------------------
# LLM: requirement detection + rating
# ---------------------------------------------------------------------------

def detect_compliance_requirement(
    comment_text: str,
    anchor_text: str,
    requirements: List[Dict],
    model: str = _DEFAULT_LOCAL_MODEL,
) -> Optional[Dict]:
    """
    Use the LLM to identify which RFP requirement a reviewer comment addresses.

    Args:
        comment_text: The reviewer's comment body.
        anchor_text: The text the comment is anchored to.
        requirements: List of requirement dicts from fetch_shredding_requirements().
        model: Ollama model name.

    Returns:
        dict or None: The matched requirement dict, or None if no match found.
    """
    if not requirements:
        return None

    # Format requirement list for LLM
    req_lines = "\n".join(
        f"[{i}] (ID:{r['id'][:8]}) Section {r.get('section','?')}: {r['source_text'][:200]}"
        for i, r in enumerate(requirements)
    )

    prompt = f"""You are reviewing a government proposal.

Reviewer comment: "{comment_text}"
Commented text: "{anchor_text}"

RFP Requirements (numbered):
{req_lines}

Which requirement number does this comment most directly address?
Reply with ONLY the number (e.g. "3") or "none" if no match applies.
"""
    response = _call_ollama(model, prompt).strip().lower()

    # Parse the index
    match = re.search(r"\d+", response)
    if match:
        idx = int(match.group())
        if 0 <= idx < len(requirements):
            return requirements[idx]

    return None


def rate_section_comment(
    section_content: str,
    comment_text: str,
    requirement_text: str,
    model: str = _DEFAULT_LOCAL_MODEL,
) -> str:
    """
    Generate an evaluation rating for the section's compliance with a requirement.

    Uses the US government five-point adjectival rating scale
    (FAR 15.305 / DFARS 215.305).

    Args:
        section_content: Current prose of the proposal section.
        comment_text: Reviewer's comment about this section.
        requirement_text: The RFP requirement being assessed.
        model: Ollama model name.

    Returns:
        str: One of 'Unacceptable', 'Marginal', 'Acceptable', 'Good', 'Outstanding'.
    """
    section_excerpt = (section_content or "")[:1200]
    prompt = f"""You are a senior government proposal evaluator using the adjectival
rating scale from FAR 15.305.

RFP Requirement:
"{requirement_text}"

Proposal section content:
\"\"\"
{section_excerpt}
\"\"\"

Reviewer comment:
"{comment_text}"

Rating scale definitions:
- Outstanding: Exceptional approach; significant strengths; no weaknesses
- Good: Thorough approach; strengths outweigh weaknesses
- Acceptable: Meets requirements; strengths and weaknesses balance
- Marginal: Deficiencies present; section needs revision to be acceptable
- Unacceptable: Fails to meet requirement; not awardable as written

Based on the requirement, the section content, and the reviewer's comment,
what is the correct adjectival rating for THIS SECTION against THIS REQUIREMENT?

Reply with ONLY one word: Outstanding, Good, Acceptable, Marginal, or Unacceptable.
"""
    raw = _call_ollama(model, prompt).strip()

    # Normalise response — find first recognised token
    for rating in RATING_SCALE:
        if rating.lower() in raw.lower():
            return rating

    # Default to Marginal when LLM gives an unparseable response
    logger.warning("Unrecognised rating response: %r — defaulting to Marginal", raw)
    return "Marginal"


def _call_ollama(model: str, prompt: str) -> str:
    """Call Ollama generate endpoint, return response text or error string."""
    try:
        from ollama_config import ollama_config  # type: ignore[import]
        result = ollama_config.generate_response(model, prompt)
        if result.get("success"):
            return result["data"].get("response", "").strip()
        return f"[LLM error: {result.get('error', 'unknown')}]"
    except Exception as exc:
        logger.exception("Ollama call failed")
        return f"[LLM error: {exc}]"


# ---------------------------------------------------------------------------
# Flat section list helper
# ---------------------------------------------------------------------------

def _flat_sections(tree: List[Dict]) -> List[Dict]:
    """
    Iteratively flatten a nested section tree into a flat list (DFS order).

    Uses an explicit stack to avoid recursion-depth limits on deep trees.
    """
    out: List[Dict] = []
    stack = list(reversed(tree))  # reversed so pop() processes first child first
    while stack:
        sec = stack.pop()
        out.append(sec)
        children = sec.get("children") or []
        if children:
            stack.extend(reversed(children))
    return out


# ---------------------------------------------------------------------------
# Main ingestion pipeline
# ---------------------------------------------------------------------------

def ingest_color_team_review(
    document_id: str,
    opportunity_id: str,
    docx_bytes: bytes,
    review_stage: str = "red",
    author_prefix: str = "",
    model: str = _DEFAULT_LOCAL_MODEL,
    job_id: Optional[str] = None,
) -> Dict:
    """
    Full color-team ingestion pipeline for a marked-up .docx file.

    Steps:
      1. Extract comments from the .docx
      2. Load proposal sections and shredding requirements
      3. For each comment:
         a. Match to a section
         b. Detect which requirement it addresses (via LLM)
         c. Rate the section's compliance (via LLM)
         d. Persist as a proposal_comment with rating metadata

    Args:
        document_id: proposal_documents.id to associate comments with.
        opportunity_id: Opportunity UUID (for shredding requirements lookup).
        docx_bytes: Raw bytes of the uploaded reviewer .docx.
        review_stage: Color team label ('pink', 'red', 'gold', 'final').
        author_prefix: Optional prefix appended to each reviewer's name.
        model: Ollama model for requirement detection and rating.
        job_id: Optional background job ID for progress reporting.

    Returns:
        dict: Summary with 'ingested', 'skipped', 'comments', 'assessment'.
    """
    def _progress(msg: str) -> None:
        """Report progress if running as a background job."""
        if job_id:
            try:
                from server.background_jobs import job_store
                job_store.update_progress(job_id, msg)
            except Exception:
                pass
        logger.info("color_team ingest [%s]: %s", document_id[:8], msg)

    # Step 1 — Extract comments
    _progress("Extracting comments from .docx")
    raw_comments = extract_docx_comments(docx_bytes)
    if not raw_comments:
        return {"ingested": 0, "skipped": 0, "comments": [], "assessment": [],
                "message": "No comments found in the uploaded document."}

    # Step 2 — Load context
    _progress("Loading document sections and requirements")
    section_tree = _build_section_tree(document_id)
    flat_sections = _flat_sections(section_tree)
    section_index = build_section_index(flat_sections)
    requirements = fetch_shredding_requirements(opportunity_id)

    ingested: List[Dict] = []
    skipped = 0
    total = len(raw_comments)

    for idx, raw in enumerate(raw_comments, 1):
        comment_text = raw["content"]
        anchor_text = raw["anchor_text"]
        author = f"{author_prefix} {raw['author']}".strip() if author_prefix else raw["author"]

        if not author:
            author = "Reviewer"

        _progress(f"Processing comment {idx}/{total}")

        # Step 3a — Match to section
        section_id = match_comment_to_section(anchor_text, comment_text, flat_sections, section_index)
        if section_id is None:
            logger.info("Could not match comment to a section; skipping: %.60s", comment_text)
            skipped += 1
            continue

        section = next((s for s in flat_sections if s["id"] == section_id), {})
        section_content = section.get("content", "")

        # Step 3b — Detect requirement
        matched_req = detect_compliance_requirement(
            comment_text, anchor_text, requirements, model
        )
        req_id = matched_req["id"] if matched_req else None
        req_text = matched_req["source_text"] if matched_req else ""

        # Step 3c — Rate compliance
        rating = "Acceptable"  # default when no requirement matched
        if matched_req:
            rating = rate_section_comment(
                section_content, comment_text, req_text, model
            )

        # Step 3d — Persist
        cmt = create_comment(
            section_id=section_id,
            author=author,
            content=comment_text,
            comment_type="color_team_review",
            review_stage=review_stage,
            anchor_text=anchor_text,
        )
        if cmt:
            # Update with rating/requirement columns
            with get_conn() as conn:
                conn.execute(
                    """UPDATE proposal_comments
                       SET rating = ?, requirement_id = ?, requirement_text = ?
                       WHERE id = ?""",
                    (rating, req_id, req_text, cmt["id"]),
                )
            cmt["rating"] = rating
            cmt["requirement_id"] = req_id
            cmt["requirement_text"] = req_text
            ingested.append(cmt)

    # Build compliance assessment per requirement
    assessment = _build_assessment(ingested, requirements)

    return {
        "ingested": len(ingested),
        "skipped": skipped,
        "comments": ingested,
        "assessment": assessment,
    }


def _build_assessment(comments: List[Dict], requirements: List[Dict]) -> List[Dict]:
    """
    Aggregate per-requirement rating from ingested comments.

    When multiple comments target the same requirement, the worst rating wins
    (conservative government evaluator stance).

    Args:
        comments: Ingested comment dicts with 'rating' and 'requirement_id'.
        requirements: Full requirement list.

    Returns:
        list[dict]: One entry per requirement that received comments, with:
            requirement_id, section, requirement_text, worst_rating,
            comment_count, comment_ids.
    """
    # rating index (lower = worse)
    rating_rank = {r: i for i, r in enumerate(RATING_SCALE)}

    by_req: Dict[str, Dict] = {}
    for cmt in comments:
        req_id = cmt.get("requirement_id")
        if not req_id:
            continue
        if req_id not in by_req:
            by_req[req_id] = {
                "requirement_id": req_id,
                "requirement_text": cmt.get("requirement_text", ""),
                "worst_rating": "Outstanding",
                "comment_count": 0,
                "comment_ids": [],
            }
        entry = by_req[req_id]
        entry["comment_count"] += 1
        entry["comment_ids"].append(cmt["id"])
        if rating_rank.get(cmt["rating"], 2) < rating_rank.get(entry["worst_rating"], 2):
            entry["worst_rating"] = cmt["rating"]

    # Annotate with section info from requirements list
    req_meta = {r["id"]: r for r in requirements}
    result = []
    for req_id, entry in by_req.items():
        meta = req_meta.get(req_id, {})
        entry["section"] = meta.get("section", "")
        entry["compliance_type"] = meta.get("compliance_type", "")
        result.append(entry)

    # Sort by worst rating (worst first)
    result.sort(key=lambda e: rating_rank.get(e["worst_rating"], 2))
    return result


# ---------------------------------------------------------------------------
# DOCX export with inline color-team comments
# ---------------------------------------------------------------------------

def export_color_team_docx(document_id: str, review_stage: Optional[str] = None) -> Optional[bytes]:
    """
    Export a proposal document with all color-team review comments embedded inline.

    Each section is rendered as a Word heading + body text, followed by a
    shaded comment block for every unresolved color-team comment on that section.

    Args:
        document_id: proposal_documents.id.
        review_stage: Optional filter (e.g. 'red') — if omitted all stages included.

    Returns:
        bytes: Raw .docx content, or None if the document was not found.

    Raises:
        ImportError: If python-docx is not installed.
    """
    from docx import Document  # type: ignore[import]
    from docx.shared import RGBColor  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]

    from proposal.drafts_db import get_document, list_comments

    doc_record = get_document(document_id, include_sections=True)
    if doc_record is None:
        return None

    word_doc = Document()

    title_para = word_doc.add_heading(
        f"{doc_record['title']} — Color Team Review", level=0
    )
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if review_stage:
        sub = word_doc.add_paragraph(f"Review stage: {review_stage.title()} Team")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    word_doc.add_paragraph()

    # Flatten section tree
    flat = _flat_sections(doc_record.get("sections") or [])

    for sec in flat:
        depth = _section_depth(sec, doc_record.get("sections") or [])
        heading_level = min(depth + 1, 4)
        number = sec.get("section_number", "").strip()
        title = sec.get("title", "").strip()
        word_doc.add_heading(f"{number}  {title}" if number else title, level=heading_level)

        content = (sec.get("content") or "").strip()
        if content:
            for line in content.splitlines():
                word_doc.add_paragraph(line) if line.strip() else word_doc.add_paragraph()

        # Fetch color-team comments for this section
        all_comments = list_comments(sec["id"], include_resolved=False)
        ct_comments = [
            c for c in all_comments
            if c.get("comment_type") == "color_team_review"
            and (review_stage is None or c.get("review_stage") == review_stage)
        ]

        for cmt in ct_comments:
            rating = cmt.get("rating", "")
            author = cmt.get("author", "Reviewer")
            stage = cmt.get("review_stage", "")
            content_text = cmt.get("content", "")
            anchor = cmt.get("anchor_text", "")

            # Add shaded comment block
            para = word_doc.add_paragraph()
            _shade_paragraph(para, "E8E8F8")  # light blue-grey

            badge = f"[{stage.upper()} | {rating}] " if rating else f"[{stage.upper()}] "
            run_badge = para.add_run(badge)
            run_badge.bold = True
            run_badge.font.color.rgb = _rating_color(rating)

            run_author = para.add_run(f"{author}: ")
            run_author.bold = True
            run_author.font.color.rgb = RGBColor(0x33, 0x33, 0x88)

            para.add_run(content_text)

            if anchor:
                anchor_para = word_doc.add_paragraph()
                _shade_paragraph(anchor_para, "F0F0F0")
                run_anchor = anchor_para.add_run(f'  \u21b3 On: "{anchor[:120]}"')
                run_anchor.italic = True
                run_anchor.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    import io
    buf = io.BytesIO()
    word_doc.save(buf)
    return buf.getvalue()


def _section_depth(section: Dict, tree: List[Dict], _depth: int = 0) -> int:
    """Find the nesting depth of a section in the tree (0 = top-level)."""
    for s in tree:
        if s["id"] == section["id"]:
            return _depth
        found = _section_depth(section, s.get("children") or [], _depth + 1)
        if found >= 0:
            return found
    return 0


def _shade_paragraph(para, hex_color: str) -> None:
    """Apply a background shading to a Word paragraph."""
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


_RATING_COLORS = {
    "Outstanding": (0x00, 0x99, 0x44),
    "Good":        (0x22, 0x66, 0xCC),
    "Acceptable":  (0x55, 0x55, 0x55),
    "Marginal":    (0xCC, 0x88, 0x00),
    "Unacceptable":(0xCC, 0x22, 0x22),
}


def _rating_color(rating: str):
    """Return an RGBColor for a given adjectival rating."""
    from docx.shared import RGBColor  # type: ignore[import]
    r, g, b = _RATING_COLORS.get(rating, (0x55, 0x55, 0x55))
    return RGBColor(r, g, b)


# ---------------------------------------------------------------------------
# Compliance assessment retrieval
# ---------------------------------------------------------------------------

def get_document_compliance_assessment(document_id: str) -> List[Dict]:
    """
    Build a compliance assessment from all stored color-team comments on a document.

    Aggregates across all sections in the document.

    Args:
        document_id: proposal_documents.id.

    Returns:
        list[dict]: Assessment entries as from _build_assessment().
    """
    with get_conn() as conn:
        rows = conn.execute(
            """SELECT pc.id, pc.rating, pc.requirement_id, pc.requirement_text,
                      pc.section_id, pc.author, pc.content, pc.review_stage
               FROM proposal_comments pc
               JOIN proposal_sections ps ON pc.section_id = ps.id
               WHERE ps.document_id = ?
                 AND pc.comment_type = 'color_team_review'
                 AND pc.rating IS NOT NULL""",
            (document_id,),
        ).fetchall()

    comments = [dict(r) for r in rows]
    if not comments:
        return []

    # We don't have a full requirements list here, so build minimal assessment
    rating_rank = {r: i for i, r in enumerate(RATING_SCALE)}
    by_req: Dict[str, Dict] = {}
    for cmt in comments:
        req_id = cmt.get("requirement_id") or "__unlinked__"
        if req_id not in by_req:
            by_req[req_id] = {
                "requirement_id": req_id,
                "requirement_text": cmt.get("requirement_text", ""),
                "worst_rating": "Outstanding",
                "comment_count": 0,
                "comment_ids": [],
                "review_stages": set(),
            }
        entry = by_req[req_id]
        entry["comment_count"] += 1
        entry["comment_ids"].append(cmt["id"])
        entry["review_stages"].add(cmt.get("review_stage", ""))
        if rating_rank.get(cmt["rating"], 2) < rating_rank.get(entry["worst_rating"], 2):
            entry["worst_rating"] = cmt["rating"]

    result = []
    for entry in by_req.values():
        entry["review_stages"] = sorted(entry["review_stages"])
        result.append(entry)

    result.sort(key=lambda e: rating_rank.get(e["worst_rating"], 2))
    return result
