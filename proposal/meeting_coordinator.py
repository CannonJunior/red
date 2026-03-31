"""
Meeting Coordinator — Agenda generation, notes recording, action item tracking.

Handles the full meeting lifecycle for proposal teams:
  - Agenda generation by meeting type (kickoff, color teams, weekly, orals)
  - Structured notes recording with action items
  - Action item tracking with due-date alerting
  - DOCX meeting summary generation

Configuration via environment variables:
    OUTPUTS_DIR    — base output path (default: outputs/proposal)
    OLLAMA_BASE_URL — Ollama endpoint for AI-assisted note summarization
    OLLAMA_MODEL    — model for summarization (default: llama3.2:3b)
"""

import logging
import os
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)


def _outputs_dir() -> Path:
    return Path(os.getenv("OUTPUTS_DIR", Path(__file__).parent.parent / "outputs" / "proposal"))


# ---------------------------------------------------------------------------
# Meeting type definitions
# ---------------------------------------------------------------------------

MEETING_TYPES = {
    "kickoff": "Proposal Kickoff",
    "pink_team": "Pink Team Review",
    "red_team": "Red Team Review",
    "gold_team": "Gold Team Review",
    "weekly_sync": "Weekly Proposal Sync",
    "orals_prep": "Orals Preparation",
    "bid_no_bid": "Bid/No-Bid Decision",
    "hotwash": "Proposal Hotwash",
    "debrief": "Government Debrief",
}

# Standard agenda templates per meeting type
AGENDA_TEMPLATES: Dict[str, List[str]] = {
    "kickoff": [
        "Welcome and introductions",
        "Opportunity overview — solicitation summary",
        "Bid/No-Bid decision confirmation",
        "Win strategy and discriminators",
        "Proposal structure and volume assignments",
        "Schedule review — due dates and color teams",
        "Team roles and responsibilities",
        "Tools, processes, and communication norms",
        "Open Q&A",
        "Action items and next steps",
    ],
    "pink_team": [
        "Review objectives and ground rules",
        "Opportunity and win strategy recap",
        "Volume-by-volume review walkthrough",
        "Compliance matrix check — are all requirements addressed?",
        "Technical approach clarity and discriminators",
        "Win themes — are they woven throughout?",
        "Reviewer scorecards and feedback",
        "Priority findings and required changes",
        "Action items, owners, and deadlines",
    ],
    "red_team": [
        "Review objectives and evaluation criteria",
        "Win strategy and discriminator check",
        "Full proposal read-through (evaluator simulation)",
        "Technical volume assessment",
        "Management volume assessment",
        "Cost reasonableness review",
        "Compliance and responsiveness check",
        "Strengths, weaknesses, and discriminators",
        "Color team scorecard review",
        "Final action items before Gold Team",
    ],
    "gold_team": [
        "Executive leadership introductions",
        "Proposal strategy and win theme confirmation",
        "Go/No-Go on final submission",
        "Price review and competitive position",
        "Executive summary review",
        "Risk acceptance and mitigation sign-off",
        "Final action items and submission checklist",
    ],
    "weekly_sync": [
        "Status round-robin — volume progress",
        "Blockers and decisions needed",
        "Schedule review — on track?",
        "Open action item review",
        "New issues and risks",
        "Assignments for next week",
    ],
    "orals_prep": [
        "Review of oral presentation requirements",
        "Presentation structure and time allocation",
        "Key messages and discriminators to emphasize",
        "Practice run (timed)",
        "Q&A simulation",
        "Feedback and improvement areas",
        "Second practice run (if time permits)",
        "Final assignments before orals",
    ],
    "bid_no_bid": [
        "Opportunity summary — solicitation overview",
        "Shipley scoring matrix walkthrough",
        "Pwin estimate and rationale",
        "Win themes and discriminators",
        "Key risks and gaps",
        "Competitive landscape",
        "Resource availability assessment",
        "Go / No-Go decision vote",
        "Next steps if Go",
    ],
    "hotwash": [
        "Outcome announcement and context",
        "Government debrief notes (if available)",
        "What went well — keep doing",
        "What to improve — stop or change",
        "Technical approach lessons",
        "Process and schedule lessons",
        "Teaming and personnel lessons",
        "Process improvement recommendations",
        "Action items for next pursuit",
    ],
    "debrief": [
        "Government evaluator introductions",
        "Overall score and ranking context",
        "Technical evaluation feedback",
        "Management evaluation feedback",
        "Past performance evaluation feedback",
        "Price position",
        "Questions for the government",
        "Internal team takeaways",
    ],
}


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class ActionItem:
    """
    One action item from a meeting.

    Attributes:
        id: Unique identifier (e.g., 'AI-001').
        action: Description of what needs to be done.
        owner: Person responsible.
        due_date: ISO date string (YYYY-MM-DD).
        status: 'open', 'in_progress', 'done', 'blocked'.
        meeting_title: Meeting where this AI was created.
        notes: Additional context.
    """
    id: str
    action: str
    owner: str
    due_date: str
    status: str = "open"
    meeting_title: str = ""
    notes: str = ""

    def is_overdue(self) -> bool:
        """Check if the action item is past its due date and not done."""
        if self.status == "done":
            return False
        try:
            return date.fromisoformat(self.due_date) < date.today()
        except ValueError:
            return False

    def days_until_due(self) -> Optional[int]:
        """Days until due date (negative = overdue)."""
        try:
            return (date.fromisoformat(self.due_date) - date.today()).days
        except ValueError:
            return None


@dataclass
class MeetingRecord:
    """
    Structured record of a completed meeting.

    Attributes:
        meeting_type: Type key (e.g., 'pink_team').
        title: Full meeting title.
        date: ISO date string.
        solicitation_number: Associated proposal.
        attendees: List of attendee names.
        agenda_items: Agenda as delivered.
        notes: Free-form meeting notes.
        decisions: List of decisions made.
        action_items: Action items generated.
        facilitator: Meeting facilitator name.
        recorder: Note-taker name.
    """
    meeting_type: str
    title: str
    date: str
    solicitation_number: str
    attendees: List[str] = field(default_factory=list)
    agenda_items: List[str] = field(default_factory=list)
    notes: str = ""
    decisions: List[str] = field(default_factory=list)
    action_items: List[ActionItem] = field(default_factory=list)
    facilitator: str = ""
    recorder: str = ""

    def to_dict(self) -> Dict:
        """Serialize to a JSON-compatible dictionary."""
        return {
            "meeting_type": self.meeting_type,
            "title": self.title,
            "date": self.date,
            "solicitation_number": self.solicitation_number,
            "attendees": self.attendees,
            "agenda_items": self.agenda_items,
            "notes": self.notes,
            "decisions": self.decisions,
            "action_items": [
                {
                    "id": ai.id, "action": ai.action, "owner": ai.owner,
                    "due_date": ai.due_date, "status": ai.status,
                    "meeting_title": ai.meeting_title, "notes": ai.notes,
                }
                for ai in self.action_items
            ],
            "facilitator": self.facilitator,
            "recorder": self.recorder,
        }


# ---------------------------------------------------------------------------
# Agenda generation
# ---------------------------------------------------------------------------

def generate_agenda(
    meeting_type: str,
    solicitation_number: str,
    proposal_title: str = "",
    meeting_date: Optional[str] = None,
    custom_items: Optional[List[str]] = None,
) -> Dict:
    """
    Generate a meeting agenda for a proposal meeting.

    Args:
        meeting_type: Meeting type key (see MEETING_TYPES).
        solicitation_number: RFP solicitation number.
        proposal_title: Proposal title for header context.
        meeting_date: ISO date string. Defaults to today.
        custom_items: Additional agenda items to append.

    Returns:
        Dict: Agenda dict with title, date, items, and context metadata.

    Raises:
        ValueError: If meeting_type is not recognized.
    """
    if meeting_type not in MEETING_TYPES:
        valid = ", ".join(sorted(MEETING_TYPES.keys()))
        raise ValueError(f"Unknown meeting_type '{meeting_type}'. Valid types: {valid}")

    items = list(AGENDA_TEMPLATES.get(meeting_type, []))
    if custom_items:
        items.extend(custom_items)

    return {
        "meeting_type": meeting_type,
        "title": MEETING_TYPES[meeting_type],
        "solicitation_number": solicitation_number,
        "proposal_title": proposal_title,
        "date": meeting_date or date.today().isoformat(),
        "agenda_items": items,
        "item_count": len(items),
    }


# ---------------------------------------------------------------------------
# Action item management
# ---------------------------------------------------------------------------

def create_action_items(
    raw_items: List[Dict],
    meeting_title: str,
    id_prefix: str = "AI",
    start_index: int = 1,
) -> List[ActionItem]:
    """
    Create ActionItem objects from raw meeting input.

    Args:
        raw_items: List of dicts with action, owner, due_date, and optional status/notes.
        meeting_title: Meeting title to associate with each action item.
        id_prefix: Prefix for auto-generated IDs.
        start_index: Starting index for ID generation.

    Returns:
        List[ActionItem]: Parsed and ID-assigned action items.
    """
    items = []
    for i, raw in enumerate(raw_items, start=start_index):
        item_id = raw.get("id") or f"{id_prefix}-{i:03d}"
        items.append(ActionItem(
            id=str(item_id),
            action=str(raw.get("action", "")),
            owner=str(raw.get("owner", "TBD")),
            due_date=str(raw.get("due_date", "")),
            status=str(raw.get("status", "open")).lower(),
            meeting_title=meeting_title,
            notes=str(raw.get("notes", "")),
        ))
    return items


def get_overdue_items(action_items: List[ActionItem]) -> List[ActionItem]:
    """
    Filter action items that are overdue (past due date and not done).

    Args:
        action_items: Full list of action items.

    Returns:
        List[ActionItem]: Overdue items sorted by due_date ascending.
    """
    overdue = [ai for ai in action_items if ai.is_overdue()]
    return sorted(overdue, key=lambda ai: ai.due_date)


def get_upcoming_items(
    action_items: List[ActionItem],
    days_ahead: int = 7,
) -> List[ActionItem]:
    """
    Filter action items due within the next N days.

    Args:
        action_items: Full list of action items.
        days_ahead: Include items due within this many days.

    Returns:
        List[ActionItem]: Upcoming items sorted by due_date ascending.
    """
    cutoff = date.today() + timedelta(days=days_ahead)
    upcoming = [
        ai for ai in action_items
        if ai.status != "done"
        and ai.days_until_due() is not None
        and 0 <= (ai.days_until_due() or 0) <= days_ahead
    ]
    return sorted(upcoming, key=lambda ai: ai.due_date)


def action_items_by_owner(action_items: List[ActionItem]) -> Dict[str, List[ActionItem]]:
    """
    Group open action items by owner.

    Args:
        action_items: Full list of action items.

    Returns:
        Dict[str, List[ActionItem]]: owner → list of open items.
    """
    grouped: Dict[str, List[ActionItem]] = {}
    for ai in action_items:
        if ai.status != "done":
            grouped.setdefault(ai.owner, []).append(ai)
    return grouped


# ---------------------------------------------------------------------------
# DOCX meeting summary
# ---------------------------------------------------------------------------

def export_meeting_summary(record: MeetingRecord, output_path: Optional[Path] = None) -> Path:
    """
    Export a meeting record as a formatted DOCX file.

    Args:
        record: Completed meeting record.
        output_path: Where to write the file. Auto-generated if None.

    Returns:
        Path: Path to the written DOCX file.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _navy = RGBColor(0x00, 0x20, 0x60)

    # Title
    heading = doc.add_heading(record.title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = _navy

    # Metadata table
    doc.add_paragraph()
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    for row_idx, (label, value) in enumerate([
        ("Solicitation",  record.solicitation_number),
        ("Date",          record.date),
        ("Facilitator",   record.facilitator or "—"),
        ("Attendees",     ", ".join(record.attendees) or "—"),
    ]):
        meta_table.cell(row_idx, 0).text = label
        meta_table.cell(row_idx, 0).paragraphs[0].runs[0].font.bold = True
        meta_table.cell(row_idx, 1).text = value

    # Agenda
    doc.add_heading("Agenda", level=1)
    for i, item in enumerate(record.agenda_items, start=1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")

    # Discussion notes
    doc.add_heading("Discussion Notes", level=1)
    doc.add_paragraph(record.notes or "[ See recording or attendee notes ]")

    # Decisions
    if record.decisions:
        doc.add_heading("Decisions Made", level=1)
        for decision in record.decisions:
            doc.add_paragraph(decision, style="List Bullet")

    # Action items
    doc.add_heading("Action Items", level=1)
    if record.action_items:
        ai_table = doc.add_table(rows=1 + len(record.action_items), cols=4)
        ai_table.style = "Table Grid"
        headers = ["ID", "Action", "Owner", "Due Date"]
        for col_idx, hdr in enumerate(headers):
            cell = ai_table.cell(0, col_idx)
            cell.text = hdr
            cell.paragraphs[0].runs[0].font.bold = True
        for row_idx, ai in enumerate(record.action_items, start=1):
            ai_table.cell(row_idx, 0).text = ai.id
            ai_table.cell(row_idx, 1).text = ai.action
            ai_table.cell(row_idx, 2).text = ai.owner
            ai_table.cell(row_idx, 3).text = ai.due_date
    else:
        doc.add_paragraph("No action items recorded.")

    # Footer
    footer_para = doc.add_paragraph()
    run = footer_para.add_run(f"Notes recorded by: {record.recorder or 'TBD'}")
    run.font.size = Pt(9)
    run.font.italic = True

    # Determine output path
    if output_path is None:
        safe_sol = record.solicitation_number.replace("/", "-").replace(" ", "_")
        safe_date = record.date.replace("-", "")
        safe_type = record.meeting_type
        out_dir = _outputs_dir() / safe_sol / "meetings"
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = out_dir / f"{safe_sol}_{safe_type}_{safe_date}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Saved meeting summary: %s", output_path)
    return output_path


# ---------------------------------------------------------------------------
# Ollama meeting summarizer
# ---------------------------------------------------------------------------

def summarize_notes(raw_notes: str, meeting_type: str = "") -> str:
    """
    Summarize raw meeting notes using Ollama.

    Falls back to returning the original notes if Ollama is unavailable.

    Args:
        raw_notes: Unstructured meeting notes text.
        meeting_type: Meeting type for context in the prompt.

    Returns:
        str: Summarized notes (or original if Ollama unavailable).
    """
    import requests as _requests

    base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434").rstrip("/")
    model = os.getenv("OLLAMA_MODEL", "llama3.2:3b")
    meeting_label = MEETING_TYPES.get(meeting_type, "proposal meeting")

    prompt = f"""Summarize the following notes from a {meeting_label}.
Extract: (1) key discussion points, (2) decisions made, (3) action items with owners.
Be concise and factual. Format as bullet points.

Notes:
{raw_notes}
"""
    try:
        resp = _requests.post(
            f"{base_url}/api/generate",
            json={"model": model, "prompt": prompt, "stream": False,
                  "options": {"num_predict": 500, "temperature": 0.2}},
            timeout=60,
        )
        resp.raise_for_status()
        return resp.json().get("response", raw_notes).strip()
    except _requests.RequestException as exc:
        logger.warning("Ollama unavailable for note summarization: %s", exc)
        return raw_notes
