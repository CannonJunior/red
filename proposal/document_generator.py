"""
Document Generator — Template-to-DOCX engine for proposal volumes.

Maps shredding requirements to proposal sections, generates prose via Ollama,
and produces formatted DOCX files from base templates.

Configuration via environment variables:
    TEMPLATES_DIR      — path to templates/proposal/ (default: auto-detected)
    OLLAMA_BASE_URL    — Ollama endpoint (default: http://localhost:11434)
    OLLAMA_MODEL       — model for prose generation (default: llama3.2:3b)
    DOC_MAX_TOKENS     — max tokens per section (default: 800)
    DOC_TEMPERATURE    — Ollama temperature (default: 0.4)
    OUTPUTS_DIR        — base output path (default: outputs/proposal)
"""

import json
import logging
import os
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional

import requests

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Configuration helpers
# ---------------------------------------------------------------------------

def _templates_dir() -> Path:
    """Resolve templates/proposal/ relative to this file's package root."""
    return Path(os.getenv("TEMPLATES_DIR", Path(__file__).parent.parent / "templates" / "proposal"))


def _outputs_dir() -> Path:
    return Path(os.getenv("OUTPUTS_DIR", Path(__file__).parent.parent / "outputs" / "proposal"))


def _ollama_url() -> str:
    base = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434").rstrip("/")
    return f"{base}/api/generate"


def _ollama_model() -> str:
    return os.getenv("OLLAMA_MODEL", "llama3.2:3b")


# ---------------------------------------------------------------------------
# Section mapping — RFP requirement categories → volume sections
# ---------------------------------------------------------------------------

# Maps shredding requirement categories to proposal volume + section heading.
# Extend by adding to SECTION_MAP or overriding via SECTION_MAP_FILE env var.
SECTION_MAP: Dict[str, List[str]] = {
    # Technical Volume targets
    "technical_requirement":    ["technical_volume", "2.0 Technical Approach"],
    "system_requirement":       ["technical_volume", "2.2 Technical Solution"],
    "interface_requirement":    ["technical_volume", "2.2 Technical Solution"],
    "security_requirement":     ["technical_volume", "2.2 Technical Solution"],
    "performance_requirement":  ["technical_volume", "2.2 Technical Solution"],
    "data_requirement":         ["technical_volume", "2.2 Technical Solution"],
    "risk_item":                ["technical_volume", "5.0 Risk Management"],

    # Management Volume targets
    "management_requirement":   ["management_volume", "1.0 Management Approach"],
    "staffing_requirement":     ["management_volume", "3.0 Key Personnel"],
    "transition_requirement":   ["management_volume", "4.0 Transition Plan"],
    "quality_requirement":      ["management_volume", "5.0 Quality Assurance Plan"],
    "subcontracting_requirement":["management_volume", "6.0 Subcontracting Plan"],

    # Past Performance targets
    "past_performance_requirement": ["past_performance", "2.0 Relevant Past Performance Reference 1"],

    # Default fallback
    "general_requirement":      ["technical_volume", "2.0 Technical Approach"],
}


def _load_section_map() -> Dict[str, List[str]]:
    """Load section map from env override file or use built-in default."""
    map_file = os.getenv("SECTION_MAP_FILE")
    if map_file and Path(map_file).is_file():
        try:
            with open(map_file) as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError) as exc:
            logger.warning("Could not load SECTION_MAP_FILE %s: %s", map_file, exc)
    return SECTION_MAP


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class Requirement:
    """
    A single requirement extracted by the shredding skill.

    Attributes:
        id: Unique identifier (e.g., REQ-001).
        category: Type of requirement (e.g., 'technical_requirement').
        text: Full requirement text from the RFP.
        reference: RFP section reference (e.g., 'SOW 3.2.1').
        priority: Compliance priority ('shall', 'should', 'may').
    """
    id: str
    category: str
    text: str
    reference: str = ""
    priority: str = "shall"


@dataclass
class SectionContent:
    """
    Generated content for one proposal section.

    Attributes:
        volume: Target volume (e.g., 'technical_volume').
        heading: Section heading text.
        requirements: Requirements mapped to this section.
        prose: LLM-generated prose for the section.
        generation_notes: Metadata about the generation.
    """
    volume: str
    heading: str
    requirements: List[Requirement] = field(default_factory=list)
    prose: str = ""
    generation_notes: str = ""


@dataclass
class DraftResult:
    """
    Result of a document generation run.

    Attributes:
        solicitation_number: RFP solicitation number.
        volumes_written: Dict of volume name → output Path.
        sections_generated: Number of sections drafted.
        requirements_addressed: Total requirements mapped.
        errors: Non-fatal errors encountered.
    """
    solicitation_number: str
    volumes_written: Dict[str, Path] = field(default_factory=dict)
    sections_generated: int = 0
    requirements_addressed: int = 0
    errors: List[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Requirement loading
# ---------------------------------------------------------------------------

def load_requirements_from_json(path: Path) -> List[Requirement]:
    """
    Load requirements from a shredding output JSON file.

    Expected format: list of objects with id, category, text, reference, priority.

    Args:
        path: Path to the JSON requirements file.

    Returns:
        List[Requirement]: Parsed requirements.

    Raises:
        FileNotFoundError: If path does not exist.
        ValueError: If file format is invalid.
    """
    if not path.is_file():
        raise FileNotFoundError(f"Requirements file not found: {path}")

    with open(path) as f:
        raw = json.load(f)

    if not isinstance(raw, list):
        raise ValueError(f"Requirements file must be a JSON array, got {type(raw).__name__}")

    reqs = []
    for i, item in enumerate(raw):
        if not isinstance(item, dict):
            logger.warning("Skipping non-dict requirement at index %d", i)
            continue
        reqs.append(Requirement(
            id=str(item.get("id", f"REQ-{i+1:03d}")),
            category=str(item.get("category", "general_requirement")).lower(),
            text=str(item.get("text", "")),
            reference=str(item.get("reference", "")),
            priority=str(item.get("priority", "shall")).lower(),
        ))

    logger.info("Loaded %d requirements from %s", len(reqs), path)
    return reqs


def load_requirements_from_csv(path: Path) -> List[Requirement]:
    """
    Load requirements from a shredding output CSV file.

    Expected columns: id, category, text, reference, priority (priority optional).

    Args:
        path: Path to the CSV requirements file.

    Returns:
        List[Requirement]: Parsed requirements.
    """
    import csv
    if not path.is_file():
        raise FileNotFoundError(f"Requirements file not found: {path}")

    reqs = []
    with open(path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader):
            req_id = row.get("id") or row.get("requirement_id") or f"REQ-{i+1:03d}"
            category = (row.get("category") or row.get("type") or "general_requirement").lower()
            text = row.get("text") or row.get("requirement") or ""
            if not text:
                continue
            reqs.append(Requirement(
                id=str(req_id),
                category=category,
                text=str(text),
                reference=str(row.get("reference", "")),
                priority=str(row.get("priority", "shall")).lower(),
            ))

    logger.info("Loaded %d requirements from CSV %s", len(reqs), path)
    return reqs


# ---------------------------------------------------------------------------
# Section mapping
# ---------------------------------------------------------------------------

def map_requirements_to_sections(
    requirements: List[Requirement],
    section_map: Optional[Dict[str, List[str]]] = None,
) -> Dict[str, SectionContent]:
    """
    Group requirements into proposal sections by category.

    Args:
        requirements: Flat list of requirements to map.
        section_map: Category → [volume, heading] mapping. Defaults to SECTION_MAP.

    Returns:
        Dict[str, SectionContent]: section_key → SectionContent.
    """
    smap = section_map or _load_section_map()
    sections: Dict[str, SectionContent] = {}

    for req in requirements:
        mapping = smap.get(req.category) or smap.get("general_requirement", ["technical_volume", "2.0 Technical Approach"])
        volume, heading = mapping[0], mapping[1]
        key = f"{volume}::{heading}"

        if key not in sections:
            sections[key] = SectionContent(volume=volume, heading=heading)
        sections[key].requirements.append(req)

    logger.info("Mapped %d requirements into %d sections", len(requirements), len(sections))
    return sections


# ---------------------------------------------------------------------------
# Prose generation via Ollama
# ---------------------------------------------------------------------------

def _build_section_prompt(section: SectionContent, company_name: str, proposal_title: str) -> str:
    """
    Build a prompt for Ollama to draft one proposal section.

    Args:
        section: Section with mapped requirements.
        company_name: Name of the proposing company.
        proposal_title: Proposal/opportunity title.

    Returns:
        str: Formatted prompt.
    """
    req_list = "\n".join(
        f"  [{r.id}] ({r.priority.upper()}) {r.text}" + (f" [Ref: {r.reference}]" if r.reference else "")
        for r in section.requirements
    )
    return f"""You are a senior GovCon proposal writer drafting a federal proposal section.

Company: {company_name}
Proposal: {proposal_title}
Section: {section.heading}

Requirements this section must address:
{req_list}

Write a professional, compliance-focused proposal section that:
1. Directly addresses each requirement listed above
2. Uses active voice and strong technical language
3. Avoids jargon without substance
4. Is appropriate for a government evaluator
5. Is 3-5 paragraphs, approximately 400-600 words

Begin the section directly — do not restate the section heading or add preamble.
"""


def generate_section_prose(
    section: SectionContent,
    company_name: str = "Our Company",
    proposal_title: str = "Proposal",
) -> str:
    """
    Generate proposal prose for one section using Ollama.

    Falls back gracefully if Ollama is unavailable.

    Args:
        section: Section with mapped requirements.
        company_name: Proposing company name for context.
        proposal_title: Proposal title for context.

    Returns:
        str: Generated prose (or placeholder if generation fails).
    """
    prompt = _build_section_prompt(section, company_name, proposal_title)
    max_tokens = int(os.getenv("DOC_MAX_TOKENS", "800"))
    temperature = float(os.getenv("DOC_TEMPERATURE", "0.4"))

    try:
        resp = requests.post(
            _ollama_url(),
            json={
                "model": _ollama_model(),
                "prompt": prompt,
                "stream": False,
                "options": {
                    "num_predict": max_tokens,
                    "temperature": temperature,
                },
            },
            timeout=120,
        )
        resp.raise_for_status()
        return resp.json().get("response", "").strip()
    except requests.RequestException as exc:
        logger.warning("Ollama unavailable for section '%s': %s", section.heading, exc)
        req_ids = ", ".join(r.id for r in section.requirements)
        return (
            f"[ DRAFT REQUIRED — Section: {section.heading} ]\n\n"
            f"This section addresses requirements: {req_ids}\n\n"
            f"Please draft content that addresses each requirement. "
            f"Refer to the requirements matrix for compliance mapping."
        )


# ---------------------------------------------------------------------------
# DOCX writing
# ---------------------------------------------------------------------------

def _load_docx_template(volume: str) -> "Document":  # type: ignore[name-defined]  # noqa: F821
    """
    Load the base DOCX template for a volume.

    Args:
        volume: Volume name (e.g., 'technical_volume').

    Returns:
        Document: Loaded python-docx Document.

    Raises:
        FileNotFoundError: If template does not exist.
    """
    from docx import Document

    template_file = _templates_dir() / f"{volume}.docx"
    if not template_file.is_file():
        raise FileNotFoundError(
            f"Template not found: {template_file}. "
            "Run: uv run templates/proposal/generate_templates.py"
        )
    return Document(str(template_file))


def _find_placeholder_paragraph(doc, placeholder: str):
    """
    Find the paragraph in a document that contains a placeholder string.

    Args:
        doc: python-docx Document.
        placeholder: Substring to search for.

    Returns:
        Paragraph or None.
    """
    for para in doc.paragraphs:
        if placeholder.upper() in para.text.upper():
            return para
    return None


def write_section_to_doc(doc, section: SectionContent) -> None:
    """
    Write a generated section into an open python-docx Document.

    Replaces the placeholder paragraph for the section, or appends if not found.

    Args:
        doc: python-docx Document.
        section: Section with generated prose.
    """
    from docx.shared import Pt

    # Try to find and replace the placeholder paragraph
    placeholder = f"[ {section.heading.upper()}"
    target_para = _find_placeholder_paragraph(doc, placeholder)

    if target_para:
        # Clear the placeholder and insert generated prose
        target_para.clear()
        target_para.add_run(section.prose)
    else:
        # Append at end
        doc.add_heading(section.heading, level=2)
        for para_text in section.prose.split("\n\n"):
            stripped = para_text.strip()
            if stripped:
                doc.add_paragraph(stripped)

    # Requirements compliance matrix note
    req_refs = " | ".join(
        f"{r.id} [{r.priority}]" + (f" {r.reference}" if r.reference else "")
        for r in section.requirements
    )
    note_para = doc.add_paragraph(f"Requirements addressed: {req_refs}")
    note_para.runs[0].font.size = Pt(8)
    note_para.runs[0].font.italic = True


# ---------------------------------------------------------------------------
# Main public API
# ---------------------------------------------------------------------------

def draft_volume(
    volume: str,
    sections: Dict[str, SectionContent],
    proposal_title: str,
    solicitation_number: str,
    company_name: str = "Our Company",
) -> Path:
    """
    Draft one proposal volume as a DOCX file.

    Args:
        volume: Volume name (e.g., 'technical_volume').
        sections: All sections keyed by '{volume}::{heading}'.
        proposal_title: Full proposal title (goes in header).
        solicitation_number: RFP solicitation number.
        company_name: Proposing company name.

    Returns:
        Path: Path to the generated DOCX file.
    """
    doc = _load_docx_template(volume)

    # Filter sections for this volume
    volume_sections = {k: s for k, s in sections.items() if s.volume == volume}
    if not volume_sections:
        logger.info("No sections for volume '%s' — writing template only", volume)

    for section in volume_sections.values():
        logger.info("Drafting section: %s", section.heading)
        if not section.prose:
            section.prose = generate_section_prose(section, company_name, proposal_title)
        write_section_to_doc(doc, section)

    # Determine output path
    safe_sol = solicitation_number.replace("/", "-").replace(" ", "_")
    today = date.today().strftime("%Y-%m-%d")
    out_dir = _outputs_dir() / safe_sol
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{safe_sol}_{volume}_{today}.docx"

    doc.save(str(out_path))
    logger.info("Saved %s: %s", volume, out_path)
    return out_path


def draft_proposal(
    requirements: List[Requirement],
    proposal_title: str,
    solicitation_number: str,
    company_name: str = "Our Company",
    volumes: Optional[List[str]] = None,
    section_map: Optional[Dict[str, List[str]]] = None,
) -> DraftResult:
    """
    Draft one or more proposal volumes from a requirements list.

    Args:
        requirements: Requirements extracted by the shredding skill.
        proposal_title: Proposal/opportunity title.
        solicitation_number: RFP solicitation number.
        company_name: Proposing company name.
        volumes: Which volumes to draft. Defaults to all unique volumes in section map.
        section_map: Category → [volume, heading] override.

    Returns:
        DraftResult: Summary including output paths and error list.
    """
    result = DraftResult(solicitation_number=solicitation_number)
    sections = map_requirements_to_sections(requirements, section_map)
    result.requirements_addressed = len(requirements)

    # Determine which volumes to draft
    all_volumes = {s.volume for s in sections.values()}
    target_volumes = [v for v in (volumes or sorted(all_volumes)) if v in all_volumes]

    if not target_volumes:
        result.errors.append("No requirements could be mapped to any volume.")
        return result

    for volume in target_volumes:
        try:
            out_path = draft_volume(
                volume, sections, proposal_title, solicitation_number, company_name,
            )
            result.volumes_written[volume] = out_path
            result.sections_generated += sum(1 for s in sections.values() if s.volume == volume)
        except FileNotFoundError as exc:
            result.errors.append(str(exc))
            logger.error("Template missing for volume '%s': %s", volume, exc)
        except Exception as exc:  # noqa: BLE001
            result.errors.append(f"Error drafting {volume}: {exc}")
            logger.error("Failed to draft volume '%s': %s", volume, exc, exc_info=True)

    return result
