"""
Hotwash — Post-proposal retrospective engine.

Facilitates structured lessons-learned sessions after proposal submission,
generates DOCX/Confluence hotwash reports, and creates improvement tickets.

Configuration via environment variables:
    OUTPUTS_DIR     — base output path (default: outputs/proposal)
    OLLAMA_BASE_URL — Ollama endpoint for AI-assisted analysis
    OLLAMA_MODEL    — model for insight generation (default: llama3.2:3b)
"""

import logging
import os
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional

import requests
from ollama_config import ollama_config

logger = logging.getLogger(__name__)


def _outputs_dir() -> Path:
    return Path(os.getenv("OUTPUTS_DIR", Path(__file__).parent.parent / "outputs" / "proposal"))


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

LESSON_CATEGORIES = [
    "technical_approach",
    "management_approach",
    "past_performance",
    "cost_price",
    "proposal_process",
    "schedule_management",
    "teaming",
    "customer_relations",
    "competitive_intelligence",
    "other",
]

IMPACT_LEVELS = ["critical", "high", "medium", "low"]


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class Lesson:
    """
    One lesson learned from a proposal debrief or internal retrospective.

    Attributes:
        id: Unique identifier (e.g., 'LL-001').
        category: Lesson category (see LESSON_CATEGORIES).
        lesson: Description of what was learned.
        impact: Impact level ('critical', 'high', 'medium', 'low').
        recommended_action: Specific process improvement to implement.
        owner: Person responsible for implementing the action.
        due_date: ISO date for when the action should be complete.
        source: Where the lesson came from ('debrief', 'internal', 'evaluator').
    """
    id: str
    category: str
    lesson: str
    impact: str = "medium"
    recommended_action: str = ""
    owner: str = "TBD"
    due_date: str = ""
    source: str = "internal"

    def to_dict(self) -> Dict:
        return {
            "id": self.id, "category": self.category, "lesson": self.lesson,
            "impact": self.impact, "recommended_action": self.recommended_action,
            "owner": self.owner, "due_date": self.due_date, "source": self.source,
        }


@dataclass
class DebriefNotes:
    """
    Notes captured from the government debrief session.

    Attributes:
        received: Whether a debrief was received.
        date: ISO date of the debrief.
        evaluators: Names or titles of government evaluators.
        overall_score: Government-assigned score (if disclosed).
        ranking: Our ranking among offerors (if disclosed).
        strengths: Cited proposal strengths.
        weaknesses: Cited proposal weaknesses.
        deficiencies: Cited deficiencies (more severe than weaknesses).
        price_position: Our price relative to awardee/others.
        raw_notes: Verbatim or paraphrased notes from the debrief.
    """
    received: bool = False
    date: str = ""
    evaluators: List[str] = field(default_factory=list)
    overall_score: str = ""
    ranking: str = ""
    strengths: List[str] = field(default_factory=list)
    weaknesses: List[str] = field(default_factory=list)
    deficiencies: List[str] = field(default_factory=list)
    price_position: str = ""
    raw_notes: str = ""


@dataclass
class HotwashRecord:
    """
    Complete hotwash record for a submitted proposal.

    Attributes:
        solicitation_number: RFP solicitation number.
        proposal_title: Full proposal title.
        outcome: 'WIN', 'LOSS', 'NO_AWARD', or 'PENDING'.
        submitted_date: ISO date of proposal submission.
        conducted_date: ISO date of hotwash session.
        facilitator: Hotwash facilitator name.
        attendees: List of attendee names.
        debrief: Government debrief notes.
        lessons: All lessons learned.
        process_improvements: High-level process changes to implement.
        action_items: Follow-up action items with owners.
        summary: Executive summary of the hotwash.
    """
    solicitation_number: str
    proposal_title: str
    outcome: str = "PENDING"
    submitted_date: str = ""
    conducted_date: str = ""
    facilitator: str = ""
    attendees: List[str] = field(default_factory=list)
    debrief: DebriefNotes = field(default_factory=DebriefNotes)
    lessons: List[Lesson] = field(default_factory=list)
    process_improvements: List[str] = field(default_factory=list)
    action_items: List[Dict] = field(default_factory=list)
    summary: str = ""

    def lessons_by_category(self) -> Dict[str, List[Lesson]]:
        """Group lessons by category."""
        grouped: Dict[str, List[Lesson]] = {}
        for lesson in self.lessons:
            grouped.setdefault(lesson.category, []).append(lesson)
        return grouped

    def lessons_by_impact(self) -> Dict[str, List[Lesson]]:
        """Group lessons by impact level (critical first)."""
        grouped: Dict[str, List[Lesson]] = {}
        for level in IMPACT_LEVELS:
            items = [l for l in self.lessons if l.impact == level]
            if items:
                grouped[level] = items
        return grouped

    def critical_lessons(self) -> List[Lesson]:
        """Return only critical-impact lessons."""
        return [l for l in self.lessons if l.impact == "critical"]

    def to_dict(self) -> Dict:
        """Serialize to a JSON-compatible dictionary."""
        return {
            "solicitation_number": self.solicitation_number,
            "proposal_title": self.proposal_title,
            "outcome": self.outcome,
            "submitted_date": self.submitted_date,
            "conducted_date": self.conducted_date,
            "facilitator": self.facilitator,
            "attendees": self.attendees,
            "debrief": {
                "received": self.debrief.received,
                "date": self.debrief.date,
                "overall_score": self.debrief.overall_score,
                "ranking": self.debrief.ranking,
                "strengths": self.debrief.strengths,
                "weaknesses": self.debrief.weaknesses,
                "deficiencies": self.debrief.deficiencies,
                "price_position": self.debrief.price_position,
            },
            "lessons": [l.to_dict() for l in self.lessons],
            "process_improvements": self.process_improvements,
            "action_items": self.action_items,
            "summary": self.summary,
        }


# ---------------------------------------------------------------------------
# Lesson builders
# ---------------------------------------------------------------------------

def create_lessons(raw_items: List[Dict], id_prefix: str = "LL") -> List[Lesson]:
    """
    Build Lesson objects from raw input dicts.

    Args:
        raw_items: List of dicts with lesson, category, impact, etc.
        id_prefix: Prefix for auto-generated IDs.

    Returns:
        List[Lesson]: Parsed lesson objects.
    """
    lessons = []
    for i, item in enumerate(raw_items, start=1):
        category = str(item.get("category", "other")).lower()
        if category not in LESSON_CATEGORIES:
            logger.warning("Unknown lesson category '%s', using 'other'", category)
            category = "other"
        impact = str(item.get("impact", "medium")).lower()
        if impact not in IMPACT_LEVELS:
            impact = "medium"

        lessons.append(Lesson(
            id=str(item.get("id", f"{id_prefix}-{i:03d}")),
            category=category,
            lesson=str(item.get("lesson", "")),
            impact=impact,
            recommended_action=str(item.get("recommended_action", "")),
            owner=str(item.get("owner", "TBD")),
            due_date=str(item.get("due_date", "")),
            source=str(item.get("source", "internal")),
        ))
    return lessons


def build_hotwash_from_dict(data: Dict) -> HotwashRecord:
    """
    Build a HotwashRecord from a plain dictionary.

    Args:
        data: Input dictionary (from API, JSON, or skill invocation).

    Returns:
        HotwashRecord: Populated record.
    """
    debrief_data = data.get("debrief", {})
    debrief = DebriefNotes(
        received=bool(debrief_data.get("received", False)),
        date=str(debrief_data.get("date", "")),
        evaluators=list(debrief_data.get("evaluators", [])),
        overall_score=str(debrief_data.get("overall_score", "")),
        ranking=str(debrief_data.get("ranking", "")),
        strengths=list(debrief_data.get("strengths", [])),
        weaknesses=list(debrief_data.get("weaknesses", [])),
        deficiencies=list(debrief_data.get("deficiencies", [])),
        price_position=str(debrief_data.get("price_position", "")),
        raw_notes=str(debrief_data.get("raw_notes", "")),
    )

    lessons = create_lessons(data.get("lessons", []))

    return HotwashRecord(
        solicitation_number=str(data.get("solicitation_number", "")),
        proposal_title=str(data.get("proposal_title", "")),
        outcome=str(data.get("outcome", "PENDING")).upper(),
        submitted_date=str(data.get("submitted_date", "")),
        conducted_date=str(data.get("conducted_date", date.today().isoformat())),
        facilitator=str(data.get("facilitator", "")),
        attendees=list(data.get("attendees", [])),
        debrief=debrief,
        lessons=lessons,
        process_improvements=list(data.get("process_improvements", [])),
        action_items=list(data.get("action_items", [])),
        summary=str(data.get("summary", "")),
    )


# ---------------------------------------------------------------------------
# AI-assisted insight generation
# ---------------------------------------------------------------------------

def generate_improvement_insights(record: HotwashRecord) -> str:
    """
    Use Ollama to synthesize cross-cutting process improvement recommendations.

    Falls back gracefully if Ollama is unavailable.

    Args:
        record: Completed hotwash record with lessons.

    Returns:
        str: Synthesized improvement recommendations, or plain summary.
    """
    base_url = ollama_config.base_url.rstrip("/")
    model = os.getenv("OLLAMA_MODEL", "llama3.2:3b")

    critical = record.critical_lessons()
    all_lessons_text = "\n".join(
        f"- [{l.category}] {l.lesson} (impact: {l.impact})"
        for l in record.lessons
    )

    prompt = f"""You are a GovCon proposal process improvement consultant.

Proposal: {record.proposal_title} ({record.solicitation_number})
Outcome: {record.outcome}
Lessons learned ({len(record.lessons)} total):
{all_lessons_text}

Synthesize the top 3-5 actionable process improvements the team should implement before
the next proposal. Focus on cross-cutting patterns, not single incidents.
Format as a numbered list with brief rationale for each.
"""
    try:
        resp = requests.post(
            f"{base_url}/api/generate",
            json={
                "model": model, "prompt": prompt, "stream": False,
                "options": {"num_predict": 600, "temperature": 0.3},
            },
            timeout=90,
        )
        resp.raise_for_status()
        return resp.json().get("response", "").strip()
    except requests.RequestException as exc:
        logger.warning("Ollama unavailable for hotwash insights: %s", exc)
        if critical:
            return "Critical areas for improvement:\n" + "\n".join(
                f"- {l.lesson}" for l in critical
            )
        return "No AI synthesis available. Review lessons above manually."


# ---------------------------------------------------------------------------
# DOCX hotwash report
# ---------------------------------------------------------------------------

def export_hotwash_docx(record: HotwashRecord, output_path: Optional[Path] = None) -> Path:
    """
    Export a hotwash report as a formatted DOCX file.

    Args:
        record: Completed HotwashRecord.
        output_path: Where to write the file. Auto-generated if None.

    Returns:
        Path: Path to the written DOCX file.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _outcome_color = {
        "WIN":      RGBColor(0x00, 0xB0, 0x50),
        "LOSS":     RGBColor(0xFF, 0x00, 0x00),
        "NO_AWARD": RGBColor(0xFF, 0xC0, 0x00),
        "PENDING":  RGBColor(0x70, 0x70, 0x70),
    }
    _navy = RGBColor(0x00, 0x20, 0x60)
    color = _outcome_color.get(record.outcome, _navy)

    # Cover
    heading = doc.add_heading("Proposal Hotwash Report", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = _navy

    doc.add_paragraph()

    # Outcome banner
    outcome_para = doc.add_paragraph()
    outcome_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = outcome_para.add_run(f"OUTCOME: {record.outcome}")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = color

    doc.add_paragraph()

    # Metadata
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = "Table Grid"
    for row_idx, (label, val) in enumerate([
        ("Solicitation",    record.solicitation_number),
        ("Proposal Title",  record.proposal_title),
        ("Submitted",       record.submitted_date or "—"),
        ("Hotwash Date",    record.conducted_date or "—"),
        ("Facilitator",     record.facilitator or "—"),
    ]):
        meta_table.cell(row_idx, 0).text = label
        meta_table.cell(row_idx, 0).paragraphs[0].runs[0].font.bold = True
        meta_table.cell(row_idx, 1).text = val

    # Executive summary
    if record.summary:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(record.summary)

    # Government debrief
    doc.add_heading("Government Debrief", level=1)
    if record.debrief.received:
        debrief_table = doc.add_table(rows=4, cols=2)
        debrief_table.style = "Table Grid"
        for row_idx, (label, val) in enumerate([
            ("Score",           record.debrief.overall_score or "Not disclosed"),
            ("Ranking",         record.debrief.ranking or "Not disclosed"),
            ("Price Position",  record.debrief.price_position or "Not disclosed"),
            ("Debrief Date",    record.debrief.date or "—"),
        ]):
            debrief_table.cell(row_idx, 0).text = label
            debrief_table.cell(row_idx, 0).paragraphs[0].runs[0].font.bold = True
            debrief_table.cell(row_idx, 1).text = val

        if record.debrief.strengths:
            doc.add_heading("Cited Strengths", level=2)
            for s in record.debrief.strengths:
                doc.add_paragraph(s, style="List Bullet")
        if record.debrief.weaknesses:
            doc.add_heading("Cited Weaknesses", level=2)
            for w in record.debrief.weaknesses:
                doc.add_paragraph(w, style="List Bullet")
        if record.debrief.deficiencies:
            doc.add_heading("Deficiencies", level=2)
            for d in record.debrief.deficiencies:
                doc.add_paragraph(d, style="List Bullet")
    else:
        doc.add_paragraph("No government debrief received or requested.")

    # Lessons learned
    doc.add_heading("Lessons Learned", level=1)
    if record.lessons:
        ll_table = doc.add_table(rows=1 + len(record.lessons), cols=5)
        ll_table.style = "Table Grid"
        for col_idx, hdr in enumerate(["ID", "Category", "Lesson", "Impact", "Recommended Action"]):
            cell = ll_table.cell(0, col_idx)
            cell.text = hdr
            cell.paragraphs[0].runs[0].font.bold = True
        for row_idx, lesson in enumerate(record.lessons, start=1):
            ll_table.cell(row_idx, 0).text = lesson.id
            ll_table.cell(row_idx, 1).text = lesson.category.replace("_", " ").title()
            ll_table.cell(row_idx, 2).text = lesson.lesson
            ll_table.cell(row_idx, 3).text = lesson.impact.title()
            ll_table.cell(row_idx, 4).text = lesson.recommended_action
    else:
        doc.add_paragraph("No lessons recorded.")

    # Process improvements
    doc.add_heading("Process Improvement Recommendations", level=1)
    if record.process_improvements:
        for item in record.process_improvements:
            doc.add_paragraph(item, style="List Number")
    else:
        doc.add_paragraph("None identified.")

    # Action items
    doc.add_heading("Follow-Up Action Items", level=1)
    if record.action_items:
        ai_table = doc.add_table(rows=1 + len(record.action_items), cols=3)
        ai_table.style = "Table Grid"
        for col_idx, hdr in enumerate(["Action", "Owner", "Due Date"]):
            cell = ai_table.cell(0, col_idx)
            cell.text = hdr
            cell.paragraphs[0].runs[0].font.bold = True
        for row_idx, ai in enumerate(record.action_items, start=1):
            ai_table.cell(row_idx, 0).text = str(ai.get("action", ""))
            ai_table.cell(row_idx, 1).text = str(ai.get("owner", "TBD"))
            ai_table.cell(row_idx, 2).text = str(ai.get("due_date", ""))
    else:
        doc.add_paragraph("No action items.")

    # Determine output path
    if output_path is None:
        safe_sol = record.solicitation_number.replace("/", "-").replace(" ", "_")
        today = date.today().strftime("%Y-%m-%d")
        out_dir = _outputs_dir() / safe_sol
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = out_dir / f"{safe_sol}_hotwash_{today}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Saved hotwash report: %s", output_path)
    return output_path
