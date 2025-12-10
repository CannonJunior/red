"""
Universal document loader for MCP tools.

Supports .txt, .pdf, .docx with caching and async loading.
"""

from pathlib import Path
from typing import List, Optional, Union, Literal
from dataclasses import dataclass
import time

# Document processing libraries
try:
    import docx
    from PyPDF2 import PdfReader
except ImportError:
    print("WARNING: Document processing libraries not installed")
    print("Install with: uv add python-docx PyPDF2")

from .config import get_config
from .errors import DocumentLoadError
from .cache import DocumentCache


@dataclass
class DocumentMetadata:
    """Metadata about a loaded document"""
    file_path: str
    file_name: str
    file_size: int
    format: str
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    char_count: int = 0
    load_time_ms: float = 0.0

    def to_dict(self):
        """Convert to dict for serialization"""
        return {
            'file_path': self.file_path,
            'file_name': self.file_name,
            'file_size': self.file_size,
            'format': self.format,
            'page_count': self.page_count,
            'word_count': self.word_count,
            'char_count': self.char_count,
            'load_time_ms': self.load_time_ms
        }


@dataclass
class LoadedDocument:
    """A loaded document with text content and metadata"""
    text: str
    metadata: DocumentMetadata

    def __len__(self):
        return len(self.text)

    def summary(self) -> str:
        """Return a summary string"""
        return (
            f"{self.metadata.file_name} "
            f"({self.metadata.format}, "
            f"{self.metadata.char_count:,} chars)"
        )

    def to_dict(self):
        """Convert to dict for serialization"""
        return {
            'text': self.text,
            'metadata': self.metadata.to_dict()
        }

    @classmethod
    def from_dict(cls, data: dict):
        """Create from dict (for cache deserialization)"""
        metadata = DocumentMetadata(**data['metadata'])
        return cls(text=data['text'], metadata=metadata)


class DocumentLoader:
    """
    Universal document loader for all MCP tools.

    Supports:
    - Single files
    - Directories (recursive)
    - Multiple formats (.txt, .pdf, .docx)
    - Caching
    - Size limits
    """

    def __init__(self, config=None, cache=None):
        """
        Initialize document loader.

        Args:
            config: MCPToolConfig instance (uses default if None)
            cache: DocumentCache instance (creates if None)
        """
        self.config = config or get_config()
        self.cache = cache or DocumentCache(
            cache_dir=self.config.cache_dir,
            enabled=self.config.enable_cache,
            ttl_hours=self.config.cache_ttl_hours
        )

    async def load(
        self,
        path: Union[str, Path],
        recursive: bool = True,
        formats: Optional[List[str]] = None
    ) -> List[LoadedDocument]:
        """
        Load document(s) from file or directory.

        Args:
            path: File path or directory path
            recursive: Recursively load from subdirectories
            formats: Allowed file formats (uses config default if None)

        Returns:
            List of LoadedDocument objects

        Raises:
            DocumentLoadError: If path invalid or loading fails
        """
        path = Path(path)
        formats = formats or self.config.supported_document_formats

        # Validate path exists
        if not path.exists():
            raise DocumentLoadError(
                error_type="path_not_found",
                message=f"Path does not exist: {path}",
                suggestions=[
                    "Check the file/folder path is correct",
                    "Ensure you have read permissions"
                ]
            )

        # Single file
        if path.is_file():
            return [await self._load_single_file(path)]

        # Directory
        elif path.is_dir():
            return await self._load_directory(path, recursive, formats)

        else:
            raise DocumentLoadError(
                error_type="invalid_path_type",
                message=f"Path is neither file nor directory: {path}"
            )

    async def _load_directory(
        self,
        dir_path: Path,
        recursive: bool,
        formats: List[str]
    ) -> List[LoadedDocument]:
        """Load all documents from directory"""

        documents = []
        total_size = 0

        # Find matching files
        pattern = "**/*" if recursive else "*"
        for file_path in dir_path.glob(pattern):
            if not file_path.is_file():
                continue

            # Check format
            if file_path.suffix.lower() not in formats:
                continue

            # Check size limit
            file_size = file_path.stat().st_size
            total_size += file_size

            if total_size > self.config.max_folder_size_mb * 1024 * 1024:
                raise DocumentLoadError(
                    error_type="folder_too_large",
                    message=f"Folder exceeds size limit: {total_size / (1024*1024):.1f}MB",
                    suggestions=[
                        f"Maximum folder size: {self.config.max_folder_size_mb}MB",
                        "Remove large files or split into multiple folders"
                    ]
                )

            # Load document
            try:
                doc = await self._load_single_file(file_path)
                documents.append(doc)
            except DocumentLoadError as e:
                # Log warning but continue with other files
                print(f"Warning: Could not load {file_path.name}: {e.message}")

        if len(documents) == 0:
            raise DocumentLoadError(
                error_type="no_documents_found",
                message=f"No documents found in {dir_path}",
                suggestions=[
                    f"Supported formats: {', '.join(formats)}",
                    f"Recursive: {recursive}"
                ]
            )

        return documents

    async def _load_single_file(self, file_path: Path) -> LoadedDocument:
        """Load a single document file"""

        start_time = time.time()

        # Check cache first
        cached = self.cache.get(str(file_path))
        if cached is not None:
            return LoadedDocument.from_dict(cached)

        # Validate file size
        file_size = file_path.stat().st_size
        max_size = self.config.max_file_size_mb * 1024 * 1024

        if file_size > max_size:
            raise DocumentLoadError(
                error_type="file_too_large",
                message=f"File exceeds size limit: {file_size / (1024*1024):.1f}MB",
                suggestions=[
                    f"Maximum file size: {self.config.max_file_size_mb}MB",
                    "Split file into smaller parts"
                ],
                context={"file_path": str(file_path)}
            )

        # Load based on format
        suffix = file_path.suffix.lower()

        try:
            if suffix == ".txt":
                text = self._load_txt(file_path)
            elif suffix == ".pdf":
                text = self._load_pdf(file_path)
            elif suffix in [".doc", ".docx"]:
                text = self._load_docx(file_path)
            else:
                raise DocumentLoadError(
                    error_type="unsupported_format",
                    message=f"Unsupported file format: {suffix}",
                    suggestions=[
                        f"Supported formats: {', '.join(self.config.supported_document_formats)}"
                    ]
                )

        except Exception as e:
            raise DocumentLoadError(
                error_type="load_failed",
                message=f"Failed to load {file_path.name}: {str(e)}",
                suggestions=[
                    "Ensure file is not corrupted",
                    "Check file permissions"
                ],
                context={"file_path": str(file_path), "error": str(e)}
            )

        # Create metadata
        load_time = (time.time() - start_time) * 1000
        metadata = DocumentMetadata(
            file_path=str(file_path),
            file_name=file_path.name,
            file_size=file_size,
            format=suffix,
            word_count=len(text.split()),
            char_count=len(text),
            load_time_ms=load_time
        )

        # Create document
        document = LoadedDocument(text=text, metadata=metadata)

        # Cache it
        self.cache.set(str(file_path), document)

        return document

    def _load_txt(self, file_path: Path) -> str:
        """Load plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _load_pdf(self, file_path: Path) -> str:
        """Load PDF file"""
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
        return "\n\n".join(text_parts)

    def _load_docx(self, file_path: Path) -> str:
        """Load Word document"""
        doc = docx.Document(file_path)
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    def combine_documents(
        self,
        documents: List[LoadedDocument],
        strategy: Literal["concatenate", "sections", "metadata"] = "sections"
    ) -> str:
        """
        Combine multiple documents into a single text.

        Args:
            documents: List of LoadedDocument objects
            strategy: How to combine
                - "concatenate": Simple join with separator
                - "sections": Add section headers per document
                - "metadata": Include file metadata in headers

        Returns:
            Combined text string
        """
        if strategy == "concatenate":
            return "\n\n".join(doc.text for doc in documents)

        elif strategy == "sections":
            parts = []
            for i, doc in enumerate(documents, 1):
                parts.append(f"--- Document {i}: {doc.metadata.file_name} ---")
                parts.append(doc.text)
                parts.append("")
            return "\n".join(parts)

        elif strategy == "metadata":
            parts = []
            for i, doc in enumerate(documents, 1):
                parts.append(f"--- Document {i} ---")
                parts.append(f"File: {doc.metadata.file_name}")
                parts.append(f"Format: {doc.metadata.format}")
                parts.append(f"Size: {doc.metadata.char_count:,} characters")
                parts.append(f"Words: {doc.metadata.word_count:,}")
                parts.append("---")
                parts.append(doc.text)
                parts.append("")
            return "\n".join(parts)

        else:
            raise ValueError(f"Unknown strategy: {strategy}")
